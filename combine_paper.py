"""Combine multiple markdown files into a single academic paper.

This script combines the individual markdown files from the reports directory
into a single cohesive paper with proper formatting, cross-references, and
image path handling. It can generate both markdown and PDF/Word outputs.

Usage:
    python combine_paper.py [--output-format FORMAT] [--output-dir DIR] [--title TITLE] [--verbose]

Output formats:
    - markdown: Single combined .md file
    - pdf: PDF using pandoc with xelatex
    - word: Microsoft Word .docx file
    - all: Generate all formats

The script hardcodes the file order and handles:
- Cross-references between sections
- Image path adjustments
- Table of contents generation
- Proper academic formatting
"""

import logging
import os
import re
import subprocess
import sys
from pathlib import Path
from typing import Dict, List, Literal, Optional

import defopt


def setup_logging(verbose: bool = False) -> logging.Logger:
    """Set up logging configuration."""
    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(
        level=level,
        format="%(asctime)s - %(levelname)s - %(message)s",
        datefmt="%H:%M:%S",
    )
    return logging.getLogger(__name__)


def load_markdown_file(file_path: Path) -> str:
    """Load a markdown file and return its content."""
    try:
        content = file_path.read_text(encoding="utf-8")
        return content
    except FileNotFoundError:
        raise FileNotFoundError(f"Markdown file not found: {file_path}")
    except Exception as e:
        raise RuntimeError(f"Error reading {file_path}: {e}")


def adjust_image_paths(content: str, base_dir: Path, output_dir: Path) -> str:
    """Adjust image paths to be relative to the output markdown file location."""
    # Pattern to match markdown image syntax: ![alt](path)
    image_pattern = r"!\[([^\]]*)\]\(([^)]+)\)"

    def replace_image_path(match):
        alt_text = match.group(1)
        image_path = match.group(2)

        # If it's a relative path from plots/, resolve it relative to base_dir
        # and then make it relative to output_dir
        if image_path.startswith("plots/"):
            # The plots are in transition_paper/plots/, resolve to output location
            full_plot_path = base_dir / image_path
            try:
                # Make relative to the output directory where pandoc will run
                rel_path = os.path.relpath(full_plot_path, output_dir)
                return f"![{alt_text}]({rel_path})"
            except ValueError:
                # If we can't make it relative, keep original
                return match.group(0)

        # If it's an absolute path, convert to relative from the output location
        if os.path.isabs(image_path):
            try:
                # Make relative to the output_dir (where pandoc will run)
                rel_path = os.path.relpath(image_path, output_dir)
                return f"![{alt_text}]({rel_path})"
            except ValueError:
                # If we can't make it relative, keep original
                return match.group(0)

        # For other relative paths, ensure they're relative to the output location
        return f"![{alt_text}]({image_path})"

    return re.sub(image_pattern, replace_image_path, content)


def format_figure_captions(content: str) -> str:
    """Move italic figure captions into the image alt text.

    Converts:
    ![old alt](path)

    *Figure X: description*

    To:
    ![Figure X: description](path)

    Pandoc uses alt text as the figure caption in both Word and PDF,
    so this avoids duplicate rendering of alt text + separate caption.
    """
    # Pattern: ![alt](path) followed by optional blank line then *Figure ...*
    # Handles multi-line captions that start with *Figure and end with *
    fig_pattern = r"!\[[^\]]*\]\(([^)]+)\)\s*\n\s*\n\*((?:Figure\s+[A-Z0-9]+[^*]*(?:\n[^*\n]*)*?))\*"

    def replace_figure(match):
        image_path = match.group(1)
        caption_text = match.group(2).strip()
        # Collapse any internal newlines in the caption to spaces
        caption_text = re.sub(r"\s*\n\s*", " ", caption_text)
        return f"![{caption_text}]({image_path})"

    return re.sub(fig_pattern, replace_figure, content, flags=re.MULTILINE | re.DOTALL)


def combine_markdown_files(
    input_files: List[Path], output_path: Path, base_dir: Path, title: str
) -> str:
    """Combine multiple markdown files into a single document."""
    logger = logging.getLogger(__name__)

    combined_content = []

    # Add YAML metadata block for proper title styling in Word/PDF
    combined_content.append("---")
    combined_content.append(f"title: \"{title}\"")
    combined_content.append("---")
    combined_content.append("")

    for file_path in input_files:
        logger.info(f"Processing {file_path.name}")

        content = load_markdown_file(file_path)

        # Adjust image paths to be relative to output directory
        output_dir = output_path.parent
        content = adjust_image_paths(content, base_dir, output_dir)

        # Format figure captions for proper Word formatting
        content = format_figure_captions(content)

        # Add content with section separator
        combined_content.append(content)
        combined_content.append("")  # Add blank line between sections

    # Join all content
    final_content = "\n".join(combined_content)

    # Write to output file
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(final_content, encoding="utf-8")

    logger.info(f"Combined markdown written to: {output_path}")
    return final_content


def convert_to_pdf(markdown_path: Path, output_path: Path) -> None:
    """Convert markdown to PDF using pandoc."""
    logger = logging.getLogger(__name__)

    # Check if pandoc is available
    try:
        subprocess.run(["pandoc", "--version"], capture_output=True, check=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        raise RuntimeError("pandoc is not installed or not available in PATH")

    # Check if xelatex is available
    try:
        subprocess.run(["xelatex", "--version"], capture_output=True, check=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        logger.warning("xelatex not found, falling back to pdflatex")
        pdf_engine = "pdflatex"
    else:
        pdf_engine = "xelatex"

    # Build pandoc command
    # Use relative paths since we'll run pandoc from the markdown directory
    markdown_rel = markdown_path.name
    output_rel = output_path.relative_to(markdown_path.parent)

    cmd = [
        "pandoc",
        markdown_rel,
        "-o",
        str(output_rel),
        f"--pdf-engine={pdf_engine}",
        "-V",
        "geometry:margin=1in",
        "-V",
        "fontsize=11pt",
        "-V",
        "documentclass=article",
        "--citeproc",  # Process citations
    ]

    # Add crossref filter if available
    try:
        subprocess.run(
            ["pandoc-crossref", "--version"], capture_output=True, check=True
        )
        cmd.extend(["--filter", "pandoc-crossref"])
        logger.info("Using pandoc-crossref for cross-references")
    except (subprocess.CalledProcessError, FileNotFoundError):
        logger.info("pandoc-crossref not available, skipping cross-references")

    try:
        logger.info(f"Converting to PDF using {pdf_engine}...")
        # Run pandoc from the directory containing the markdown file
        # This ensures that relative image paths are resolved correctly
        result = subprocess.run(
            cmd, capture_output=True, text=True, check=True, cwd=markdown_path.parent
        )
        logger.info(f"PDF generated: {output_path}")
    except subprocess.CalledProcessError as e:
        logger.error(f"PDF conversion failed: {e}")
        logger.error(f"Pandoc stderr: {e.stderr}")
        raise RuntimeError(f"PDF conversion failed: {e.stderr}")


def _create_word_reference(output_dir: Path) -> Optional[Path]:
    """Create a reference Word document with page numbers and current date in footer."""
    logger = logging.getLogger(__name__)
    try:
        from datetime import datetime

        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        doc = Document()

        # Add footer with page number and date to all sections
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.alignment = 1  # Center alignment

            # Add current date
            current_date = datetime.now().strftime("%B %d, %Y")
            run = paragraph.add_run(f"{current_date}    |    Page ")

            # Add auto-updating page number field
            fld_char_begin = OxmlElement("w:fldChar")
            fld_char_begin.set(qn("w:fldCharType"), "begin")
            run2 = paragraph.add_run()
            run2._element.append(fld_char_begin)

            instr_text = OxmlElement("w:instrText")
            instr_text.text = "PAGE"
            run3 = paragraph.add_run()
            run3._element.append(instr_text)

            fld_char_end = OxmlElement("w:fldChar")
            fld_char_end.set(qn("w:fldCharType"), "end")
            run4 = paragraph.add_run()
            run4._element.append(fld_char_end)

        ref_path = output_dir / "_reference.docx"
        doc.save(str(ref_path))
        logger.info(f"Created Word reference template: {ref_path}")
        return ref_path

    except ImportError:
        logger.warning("python-docx not available, skipping Word footer customization")
        return None
    except Exception as e:
        logger.warning(f"Failed to create Word reference template: {e}")
        return None


def convert_to_word(markdown_path: Path, output_path: Path) -> None:
    """Convert markdown to Word document using pandoc."""
    logger = logging.getLogger(__name__)

    # Check if pandoc is available
    try:
        subprocess.run(["pandoc", "--version"], capture_output=True, check=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        raise RuntimeError("pandoc is not installed or not available in PATH")

    # Build pandoc command for Word
    # Use relative paths since we'll run pandoc from the markdown directory
    markdown_rel = markdown_path.name
    output_rel = output_path.relative_to(markdown_path.parent)

    # Create a reference docx with page numbers and date in footer
    reference_docx = _create_word_reference(markdown_path.parent)

    cmd = [
        "pandoc",
        markdown_rel,
        "-o",
        str(output_rel),
        "--citeproc",  # Process citations
    ]

    if reference_docx:
        # Path must be relative to markdown_path.parent since pandoc runs from there
        ref_rel = reference_docx.relative_to(markdown_path.parent)
        cmd.extend(["--reference-doc", str(ref_rel)])

    try:
        logger.info("Converting to Word document...")
        # Run pandoc from the directory containing the markdown file
        # This ensures that relative image paths are resolved correctly
        result = subprocess.run(
            cmd, capture_output=True, text=True, check=True, cwd=markdown_path.parent
        )
        logger.info(f"Word document generated: {output_path}")
    except subprocess.CalledProcessError as e:
        logger.error(f"Word conversion failed: {e}")
        logger.error(f"Pandoc stderr: {e.stderr}")
        raise RuntimeError(f"Word conversion failed: {e.stderr}")
    finally:
        # Clean up temporary reference document
        if reference_docx and reference_docx.exists():
            reference_docx.unlink()


def main(
    *,
    output_format: Literal["markdown", "pdf", "word", "all"] = "all",
    output_dir: str = "transition_paper",
    title: str = "When Rigid Blocs Crack: Elite-Coordinated Voter Switching in an Identity-Based Party System",
    verbose: bool = False,
) -> int:
    """Combine markdown files into a single academic paper.

    Args:
        output_format: Output format ('markdown', 'pdf', 'word', or 'all')
        output_dir: Directory for output files
        title: Title of the paper
        verbose: Enable verbose logging

    Returns:
        Exit code (0 for success, 1 for failure)
    """
    logger = setup_logging(verbose)

    # Define input files in order
    reports_dir = Path("transition_paper")
    input_files = [
        reports_dir / "00_abstract.md",
        reports_dir / "01_intro.md",
        reports_dir / "01b_puzzle.md",
        reports_dir / "01c_framework.md",
        reports_dir / "02_methods.md",
        reports_dir / "03_results.md",
        reports_dir / "04_conclusions.md",
        reports_dir / "05_endmatter.md",
        reports_dir / "10_references.md",
        reports_dir / "09_appendix.md",
    ]

    # Verify all input files exist
    missing_files = [f for f in input_files if not f.exists()]
    if missing_files:
        logger.error(f"Missing input files: {[f.name for f in missing_files]}")
        return 1

    # Set up output directory
    output_dir_path = Path(output_dir)
    output_dir_path.mkdir(parents=True, exist_ok=True)

    # Generate base filename
    base_filename = "paper"

    try:
        # Always generate markdown first
        markdown_path = output_dir_path / f"{base_filename}.md"
        logger.info("Combining markdown files...")

        combine_markdown_files(input_files, markdown_path, reports_dir, title)

        # Generate additional formats based on request
        if output_format in ["pdf", "all"]:
            pdf_path = output_dir_path / f"{base_filename}.pdf"
            convert_to_pdf(markdown_path, pdf_path)

        if output_format in ["word", "all"]:
            word_path = output_dir_path / f"{base_filename}.docx"
            convert_to_word(markdown_path, word_path)

        logger.info("Paper generation completed successfully!")
        return 0

    except Exception as e:
        logger.error(f"Paper generation failed: {e}")
        return 1


if __name__ == "__main__":
    defopt.run(
        main,
        short={
            "output-format": "f",
            "output-dir": "d",
            "title": "t",
            "verbose": "v",
        },
    )
